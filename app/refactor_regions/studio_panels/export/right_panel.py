import streamlit as st
from docx import Document
from io import BytesIO


def render_right_panel(col):
    with col:
        st.header("📦 Export Outputs", divider="gray")

        # --------------------------------------------------------------------
        # Title + Subtitle – metadata (NO rewriting the widget-owned keys)
        # --------------------------------------------------------------------
        title = st.text_input(
            "Document Title",
            value=st.session_state.get("export_title_input", "Untitled Document"),
            key="export_title_input",
        )

        subtitle = st.text_input(
            "Subtitle",
            value=st.session_state.get("export_subtitle_input", ""),
            key="export_subtitle_input",
        )

        st.markdown("---")
        st.subheader("Download Files", divider="gray")

        # --------------------------------------------------------------------
        # Guardrails: verify the export is ready
        # --------------------------------------------------------------------
        md = st.session_state.get("export_markdown", "")
        html = st.session_state.get("export_html", "")

        if not md or not html:
            st.info("No export data available yet. Generate a document in the center panel.")
            return

        # --------------------------------------------------------------------
        # HTML Download
        # --------------------------------------------------------------------
        st.download_button(
            label="⬇️ Download HTML",
            data=html.encode("utf-8"),
            file_name=f"{title.replace(' ', '_')}.html",
            mime="text/html",
            use_container_width=True,
        )

        # --------------------------------------------------------------------
        # Markdown Download
        # --------------------------------------------------------------------
        st.download_button(
            label="⬇️ Download Markdown",
            data=md.encode("utf-8"),
            file_name=f"{title.replace(' ', '_')}.md",
            mime="text/markdown",
            use_container_width=True,
        )

        # --------------------------------------------------------------------
        # DOCX Download (generated dynamically)
        # --------------------------------------------------------------------
        docx_buffer = build_docx(title, subtitle, st.session_state.get("draft_text", ""), md)
        st.download_button(
            label="⬇️ Download DOCX",
            data=docx_buffer,
            file_name=f"{title.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.caption("Export tab — right panel controls (2025 modular architecture)")


# ======================================================================
# Build DOCX from markdown + metadata
# ======================================================================
def build_docx(title, subtitle, final_text, markdown_text):
    """
    Convert the exported content into a .docx binary stream.
    """
    doc = Document()

    # Title
    doc.add_heading(title, level=0)

    # Subtitle
    if subtitle:
        doc.add_paragraph(subtitle)

    doc.add_paragraph("")  # spacing

    # The final text body
    doc.add_heading("Final Draft", level=1)
    doc.add_paragraph(final_text)

    # Raw Markdown block (optional but extremely useful)
    doc.add_heading("Markdown Source", level=1)
    doc.add_paragraph(markdown_text)

    # Write to memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
