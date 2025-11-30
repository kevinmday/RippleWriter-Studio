from docx import Document

def build_docx_file(title, subtitle, sections):
    document = Document()

    # Title
    document.add_heading(title, level=1)

    # Subtitle
    if subtitle:
        document.add_heading(subtitle, level=2)

    # Sections
    for header, text in sections:
        document.add_heading(header, level=2)

        for line in (text or "").split("\n"):
            if line.strip():
                document.add_paragraph(line)

    return document
